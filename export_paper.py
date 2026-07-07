from docx import Document
from pathlib import Path

src = Path(r'c:\Users\234052\Desktop\talk-English\研究論文_草稿v6_R4SA16_瀧下敦斗_完成版.md')
out = Path(r'c:\Users\234052\Desktop\talk-English\研究論文_草稿v6_R4SA16_瀧下敦斗_完成版.docx')

doc = Document()
doc.add_heading('AIを活用した会話型英語学習アプリの設計と実装', level=1)
text = src.read_text(encoding='utf-8')
for line in text.splitlines():
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('[') and ']' in line:
        doc.add_paragraph(line)
    elif line.strip():
        doc.add_paragraph(line)
    else:
        doc.add_paragraph()
doc.save(out)
print(out)
